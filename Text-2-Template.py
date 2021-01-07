import time
from tkinter import *
from tkinter import filedialog as fd
from tkinter.ttk import Progressbar

import docx
import os


def openfile():
    global path
    username = os.getlogin()  # Fetch username
    path = fd.askopenfilename(initialdir=f'//Users//{username}//Desktop', title="Select word file",
                              filetypes=(("all files", "*.*"), ("word files", "*.docx")))
    process()


def bar():
    for x in range(10):
        progress['value'] += 10
        root.update_idletasks()
        time.sleep(0.05)
    complete.place(y=320, x=130)
    root.update_idletasks()


def process():
    bar()
    retirement = False
    general = False
    starting = False
    resiliency = False
    username = os.getlogin()  # Fetch username
    doc = docx.Document(path)
    actnum = doc.tables[0].cell(0, 5).text
    file = open(f'C:\\Users\\{username}\\Desktop\\Activity{actnum}.txt', 'w')
    file.write("CLINICAL RECORD\n\n")
    i = 60
    inc = -10
    while True:
        file.write(doc.paragraphs[i].text + '\n')
        if "PROGRESS NOTES" in doc.paragraphs[i].text:
            file.write('--------------------------------------------------------------------------------\n')
        i += 1
        inc += 1
        if "GENERAL ISSUES" in doc.paragraphs[i].text:
            break
    # DECISION TREE
    if "Select an option" not in doc.paragraphs[i + 2].text or doc.paragraphs[i + 3].text != "Other:":
        general = True
        while True:
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
            if "STARTING OUT" in doc.paragraphs[i].text:
                break
    elif "Select an option" not in doc.paragraphs[i + 12].text or doc.paragraphs[i + 13].text != "Other:":
        starting = True
        i += 10
        for x in range(4):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif "Select an option" not in doc.paragraphs[i + 16].text or doc.paragraphs[i + 7].text != "Other:":
        retirement = True
        i += 14
        for x in range(4):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif "Select an option" not in doc.paragraphs[i + 20].text or doc.paragraphs[i + 21].text != "Other:":
        resiliency = True
        i += 18
        for x in range(4):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    i = 98 + inc
    file.write("\n\n")
    # NEXT SECTION
    if general:
        while True:
            i += 1
            if "GENERAL ISSUES" in doc.paragraphs[i].text:
                break
        for x in range(27):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif starting:
        while True:
            i += 1
            if "STARTING OUT" in doc.paragraphs[i].text:
                break
        for x in range(4):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif retirement:
        while True:
            i += 1
            if "RETIREMENT PLANNING" in doc.paragraphs[i].text:
                break
        for x in range(4):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif resiliency:
        while True:
            i += 1
            if "RESILIENCY COACHING" in doc.paragraphs[i].text:
                break
        for x in range(6):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    file.write('\n\n')
    # TOP LINE
    while True:
        i += 1
        if "Next steps/homework" in doc.paragraphs[i].text:
            file.write(doc.paragraphs[i].text + '\n')
            break
    # NEXT SECTION
    if general:
        while True:
            i += 1
            if "GENERAL ISSUES" in doc.paragraphs[i].text:
                break
        for x in range(10):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif starting:
        while True:
            i += 1
            if "STARTING OUT" in doc.paragraphs[i].text:
                break
        for x in range(10):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif retirement:
        while True:
            i += 1
            if "RETIREMENT PLANNING" in doc.paragraphs[i].text:
                break
        for x in range(5):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    elif resiliency:
        while True:
            i += 1
            if "RESILIENCY COACHING" in doc.paragraphs[i].text:
                break
        for x in range(6):
            file.write(doc.paragraphs[i].text + '\n')
            i += 1
    file.write('\n\n')
    while True:
        i += 1
        if "What could help the next" in doc.paragraphs[i].text:
            break
    while True:
        file.write(doc.paragraphs[i].text + '\n')
        i += 1
        if "ession #" in doc.paragraphs[i].text:
            break
    file.write('\n\n--------------------------------------------------------------------------------\n\n')
    # TEST SESSIONS
    for q in range(6):
        if "[X]" in doc.paragraphs[i + 1].text:
            while True:
                file.write(doc.paragraphs[i].text + '\n')
                i += 1
                if "D. Have I" in doc.paragraphs[i].text:
                    while True:
                        i += 1
                        if "What could help" in doc.paragraphs[i].text:
                            break
                if "ession #" in doc.paragraphs[i].text:
                    file.write('\n\n'
                               '--------------------------------------------------------------------------------\n\n')
                    break
                elif "FILE CLOSURE" in doc.paragraphs[i].text:
                    break
        else:
            while True:
                i += 1
                if "ession #" in doc.paragraphs[i].text:
                    break
                elif "FILE CLOSURE" in doc.paragraphs[i].text:
                    break
    file.write(' \n--------------------------------------------------------------------------------\n\n')
    # LAST SECTION
    while True:
        file.write(doc.paragraphs[i].text + '\n')
        if "Preferred Provider" in doc.paragraphs[i].text:
            break
        i += 1
        if "Clinical File Sub" in doc.paragraphs[i].text:
            while True:
                i += 1
                if "Counsellor ID" in doc.paragraphs[i].text:
                    break
    file.close()


root = Tk()
root.title("Morneau Shepell Notes to Plain Text")
root.geometry("400x400")
bg = Label(root, bg="#252627")
bg.place(relheight=1, relwidth=1)
complete = Label(root, bg="#252627", fg="green", font="Montserrat 20 italic", text="Complete!")

title = Label(root, bg="#D3D4D9", fg="#252627", font="Montserrat 28 bold italic", text="Notes to Plain Text")
title.pack(pady=0, ipady=10, ipadx=100)

progress = Progressbar(root, orient=HORIZONTAL,
                       length=300, mode='determinate')
progress.place(x=50, y=300)

button = Button(root, text="Select File", command=openfile, font="Montserrat 20 bold", bg="#BB0A21", fg="#FFF9FB")
button.pack(pady=100)
root.mainloop()
